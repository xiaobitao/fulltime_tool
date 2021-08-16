from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Pt
from docx.oxml.ns import qn
from datetime import datetime

from util.drawing import draw_waterfall, draw_a_serias, draw_b_serias, draw_freq, draw_fluctuation
from util.rest import InitReport
from util.wimlog import error,info, debug
from util.config import get_config
from util.util import get_current_directory

class GenReport():
    def __init__(self):
        # init param
        self.document = Document()
        self.document.styles['Normal'].font.name = u'宋体'
        self.document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        self.document.styles['TOCHeading'].font.name = u'宋体'
        self.document.styles['TOCHeading']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        self.document.styles['TOCHeading'].font.size = Pt(6)

        self.document.styles['Heading 1'].font.name = u'宋体'
        self.document.styles['Heading 1']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        self.document.styles['Heading 1'].font.size = Pt(18)

        print(self.document.styles)
        self.tmpdir = "./tmp/"
        self.config = get_config()


    def gen_report(self, filename=""):
        rest_client = InitReport()
        self.data = rest_client.get_init_report()
        print(self.data)
        if self.data is None:
            error("Init report rest api failed")
            return ""
        self._gen_title()
        self._gen_first_paragraph()
        self._gen_sec_paragraph()
        self._gen_third_paragraph()
        self._gen_four_paragraph()
        now = datetime.now()
        strtime = now.strftime("%Y%m%d_%H_%M_%S")
        if len(filename) == 0:
            filename = get_current_directory() + "/初始化报告_%s.docx" % strtime
        self._save_report(filename=filename)
        return filename

    def _gen_title(self):
        p = self.document.add_heading(u"武汉地铁7号线仪表评估报表", level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1 = self.document.add_paragraph(u"————	2019年11月5日星期二")
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def _gen_first_paragraph(self):

        # Set a cell background (shading) color to RGB D9D9D9. 
        
        # cell._tc.get_or_add_tcPr().append(shading_elm)

        p = self.document.add_heading(u"1.仪表信息", level=1)
        # p.stype = self.document.styles['Heading 1']
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self.document.add_paragraph(
            u"武汉地铁7号线湖工站-野芷湖站光纤传感结构监测系统20xx年x月xx日仪表信息")
        table = self.document.add_table(rows=3, cols=5)
        table.style = 'TableGrid'
        table.cell(0, 0).text = u"仪表号"
        table.cell(0, 1).text = u"采集频率"
        table.cell(0, 2).text = u"通道号"
        table.cell(0, 3).text = u"传感器类型"
        table.cell(0, 4).text = u"光栅总个数"
        for i in range(0, 5):
            shading_elm = parse_xml(r'<w:shd {} w:fill="0090FF"/>'.format(nsdecls('w')))
            table.cell(0, i)._tc.get_or_add_tcPr().append(shading_elm)

        table.cell(1, 0).text = u"0001"
        table.cell(1, 1).text = u"1Khz"

        table.cell(1, 2).text = u"通道0"
        table.cell(2, 2).text = u"通道1"

        table.cell(1, 3).text = u"光纤传感器"
        table.cell(2, 3).text = u"光纤传感器"

        table.cell(1, 4).text = u"530"
        table.cell(2, 4).text = u"640"

        table.cell(1, 1).merge(table.cell(2,1))
        table.cell(1, 0).merge(table.cell(2,0))

        p2 = self.document.add_paragraph("表1-1 仪表主要信息")
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.style = self.document.styles["TOCHeading"]

    def _gen_sec_paragraph(self):
        #TODO Many should write in config file
        p = self.document.add_heading(u"2.整体评估", level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        table = self.document.add_table(rows=5, cols=4)
        table.style = 'TableGrid'

        table.cell(0, 0).text = u"通道号"
        table.cell(0, 1).text = u"区域"
        table.cell(0, 2).text = u"光栅总个数"
        table.cell(0, 3).text = u"有效光栅点占比"

        for i in range(0, 4):
            shading_elm = parse_xml(r'<w:shd {} w:fill="0090FF"/>'.format(nsdecls('w')))
            table.cell(0, i)._tc.get_or_add_tcPr().append(shading_elm)

        table.cell(1, 0).text = u"通道0"
        table.cell(1, 0).merge(table.cell(2, 0))
        table.cell(3, 0).text = u"通道1"
        table.cell(3, 0).merge(table.cell(4, 0))

        table.cell(1, 1).text = u"非减震区"
        table.cell(2, 1).text = u"减震区"
        table.cell(3, 1).text = u"非减震区"
        table.cell(4, 1).text = u"减震区"

        table.cell(1, 2).text = u"410"
        table.cell(2, 2).text = u"120"
        table.cell(3, 2).text = u"120"
        table.cell(4, 2).text = u"422"

        table.cell(1, 3).text = u"0.9"
        table.cell(2, 3).text = u"0.9"
        table.cell(3, 3).text = u"0.9"
        table.cell(4, 3).text = u"0.9"

        return True

    def get_msg(self, key, value1, value2=0):
        if key not in self.config:
            return u"没有配置说明"
        sect = self.config[key]
        good_val = float(sect["good"])
        super_val = float(sect["super"])
        normal_val = float(sect["normal"])
        if abs(value1) < super_val:
            return sect["super_msg"]
        elif abs(value1) < good_val:
            return sect["good_msg"]
        elif abs(value1) < normal_val:
            return sect["normal_msg"]
        return "没有说明"


    def _gen_third_paragraph(self):
        p = self.document.add_heading(u"3.局部评估", level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # This may repeat 
        p1 = self.document.add_heading(u"3.1区域1", level=2)
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT

        table = self.document.add_table(rows=31, cols=5)
        table.style = 'TableGrid'

        

        table.cell(0, 0).text = u"指标名称"
        table.cell(0, 1).text = u"时间段"
        table.cell(0, 2).text = u"计算结果"
        table.cell(0, 3).text = u"状态备注"
        table.cell(0, 4).text = u"状态稳定性"

        for i in range(0, 5):
            shading_elm = parse_xml(r'<w:shd {} w:fill="0090FF"/>'.format(nsdecls('w')))
            table.cell(0, i)._tc.get_or_add_tcPr().append(shading_elm)
        i = 1
        table.cell(i, 0).text = u"有效采集点比"
        table.cell(i, 1).text = u"——"
        table.cell(i, 2).text = u"0.9"
        table.cell(i, 3).text = u""
        table.cell(i, 4).text = u"——"
        
        i += 1
        #数据波动范围
        midnight_amp = self.data["vibraAmp0"]
        table.cell(i, 0).text = u"数据波动范围"
        table.cell(i, 1).text = u"凌晨"
        table.cell(i, 2).text = u"{:.1f}~{:.1f}".format(midnight_amp[0], midnight_amp[1])
        table.cell(i, 3).text = u""
        table.cell(i, 4).text = u"优"
        
        vech_amp = self.data["vibraAmp2"]
        # table.cell(3, 0).text = u"有效采集点比"
        table.cell(i+1, 1).text = u"过车"
        table.cell(i+1, 2).text = u"{:.1f}~{:.1f}".format(vech_amp[0], vech_amp[1])
        table.cell(i+1, 3).text = u""
        table.cell(i+1, 4).text = u"中"

        table.cell(i, 0).merge(table.cell(i+1, 0))

        i += 2
        #杂波占比
        mid_noise = self.data["noiseProp0"]
        table.cell(i, 0).text = u"杂波占比"
        table.cell(i, 1).text = u"凌晨"
        table.cell(i, 2).text = u"{:.1f}".format(mid_noise)
        table.cell(i, 3).text = self.get_msg("init_noise", mid_noise)
        table.cell(i, 4).text = u"优"

        vech_noise = self.data["noiseProp2"]
        table.cell(i+1, 1).text = u"过车"
        table.cell(i+1, 2).text = u"{:.1f}".format(vech_noise)
        table.cell(i+1, 3).text = self.get_msg("init_noise", vech_noise)
        table.cell(i+1, 4).text = u"中"

        table.cell(i, 0).merge(table.cell(i+1, 0))

        i += 2
        #绝对值平均值
        mid_means = self.data["meanAbs0"]
        table.cell(i, 0).text = u"绝对值平均值"
        table.cell(i, 1).text = u"凌晨"
        table.cell(i, 2).text = u"{:.1f}".format(mid_means)
        table.cell(i, 3).text = self.get_msg("init_meanabs", mid_means)
        table.cell(i, 4).text = u"优"

        vech_means = self.data["meanAbs2"]
        table.cell(i+1, 1).text = u"过车"
        table.cell(i+1, 2).text = u"{:.1f}".format(vech_means)
        table.cell(i+1, 3).text = self.get_msg("init_meanabs", vech_means)
        table.cell(i+1, 4).text = u"中"

        table.cell(i, 0).merge(table.cell(i+1, 0))

        i += 2
        # 数据波动方差
        mid_var = self.data["var0"]
        table.cell(i, 0).text = u"数据波动方差"
        table.cell(i, 1).text = u"凌晨"
        table.cell(i, 2).text = u"{:.1f}".format(mid_var)
        table.cell(i, 3).text = self.get_msg("init_vibra", mid_var)
        table.cell(i, 4).text = u"优"

        vech_var = self.data["var2"]
        table.cell(i+1, 1).text = u"过车"
        table.cell(i+1, 2).text = u"{:.1f}".format(vech_var)
        table.cell(i+1, 3).text = self.get_msg("init_vibra", vech_var)
        table.cell(i+1, 4).text = u"中"

        table.cell(i, 0).merge(table.cell(i+1, 0))

        i+=2
        # 信息熵
        mid_entropy = self.data["entropy0"]
        table.cell(i, 0).text = u"信息熵"
        table.cell(i, 1).text = u"凌晨"
        table.cell(i, 2).text =  u"{:.1f}".format(mid_entropy)
        table.cell(i, 3).text = self.get_msg("init_entropy", mid_entropy)
        table.cell(i, 4).text = u"优"
        vech_entropy = self.data["entropy2"]
        table.cell(i+1, 1).text = u"过车"
        table.cell(i+1, 2).text = self.get_msg("init_entropy", vech_entropy)
        table.cell(i+1, 3).text = u""
        table.cell(i+1, 4).text = u"中"

        table.cell(i, 0).merge(table.cell(i+1, 0))

        i+=2
        # 过零率
        mid_crosszero = self.data["zeroCrossRate0"]
        table.cell(i, 0).text = u"过零率"
        table.cell(i, 1).text = u"凌晨"
        table.cell(i, 2).text = u"{:.1f}".format(mid_crosszero)
        table.cell(i, 3).text = self.get_msg("init_zerocross", mid_crosszero)
        table.cell(i, 4).text = u"优"

        vech_crosszero = self.data["zeroCrossRate2"]
        table.cell(i+1, 1).text = u"过车"
        table.cell(i+1, 2).text = u"{:.1f}".format(vech_crosszero)
        table.cell(i+1, 3).text = self.get_msg("init_zerocross", vech_crosszero)
        table.cell(i+1, 4).text = u"中"

        table.cell(i, 0).merge(table.cell(i+1, 0))

        i+=2
        # 峰值系数
        mid_peak = self.data["peakFactor0"]
        table.cell(i, 0).text = u"峰值系数"
        table.cell(i, 1).text = u"凌晨"
        table.cell(i, 2).text = u"{:.1f}".format(mid_peak)
        table.cell(i, 3).text = u""
        table.cell(i, 4).text = u"优"

        vech_peak = self.data["peakFactor2"]
        table.cell(i+1, 1).text = u"过车"
        table.cell(i+1, 2).text = u"{:.1f}".format(vech_peak)
        table.cell(i+1, 3).text = u""
        table.cell(i+1, 4).text = u"中"

        table.cell(i, 0).merge(table.cell(i+1, 0))

        i+=2
        # 形状参数
        mid_shape = self.data["shapeParam0"]
        table.cell(i, 0).text = u"形状参数"
        table.cell(i, 1).text = u"凌晨"
        table.cell(i, 2).text = u"{:.1f}".format(mid_shape)
        table.cell(i, 3).text = u""
        table.cell(i, 4).text = u"优"

        vech_shape = self.data["shapeParam2"]
        table.cell(i+1, 1).text = u"过车"
        table.cell(i+1, 2).text = u"{:.1f}".format(vech_shape)
        table.cell(i+1, 3).text = u""
        table.cell(i+1, 4).text = u"中"

        table.cell(i, 0).merge(table.cell(i+1, 0))

        i+=2
        # 偏度参数
        mid_skew = self.data["skewParam0"]
        table.cell(i, 0).text = u"偏度参数"
        table.cell(i, 1).text = u"凌晨"
        table.cell(i, 2).text = u"{:.1f}".format(mid_skew)
        table.cell(i, 3).text = u""
        table.cell(i, 4).text = u"优"

        vech_skew = self.data["skewParam2"]
        table.cell(i+1, 1).text = u"过车"
        table.cell(i+1, 2).text =  u"{:.1f}".format(vech_skew)
        table.cell(i+1, 3).text = u""
        table.cell(i+1, 4).text = u"中"

        table.cell(i, 0).merge(table.cell(i+1, 0))

        i+=2
        # 脉冲因子
        mid_pluse = self.data["pulseFactor0"]
        table.cell(i, 0).text = u"脉冲因子"
        table.cell(i, 1).text = u"凌晨"
        table.cell(i, 2).text = u"{:.1f}".format(mid_pluse)
        table.cell(i, 3).text = u""
        table.cell(i, 4).text = u"优"

        vech_pluse = self.data["pulseFactor2"]
        table.cell(i+1, 1).text = u"过车"
        table.cell(i+1, 2).text = u"{:.1f}".format(vech_pluse)
        table.cell(i+1, 3).text = u""
        table.cell(i+1, 4).text = u"中"

        table.cell(i, 0).merge(table.cell(i+1, 0))

        i+=2
        # 裕度因子
        mid_margin = self.data["marginFactor0"]
        table.cell(i, 0).text = u"裕度因子"
        table.cell(i, 1).text = u"凌晨"
        table.cell(i, 2).text = u"{:.1f}".format(mid_margin)
        table.cell(i, 3).text = u""
        table.cell(i, 4).text = u"优"

        vech_margin = self.data["marginFactor2"]
        table.cell(i+1, 1).text = u"过车"
        table.cell(i+1, 2).text = u"{:.1f}".format(vech_margin)
        table.cell(i+1, 3).text = u""
        table.cell(i+1, 4).text = u"中"

        table.cell(i, 0).merge(table.cell(i+1, 0))

        
        i+=2
        # 峭度

        table.cell(i, 0).text = u"峭度"
        table.cell(i, 1).text = u"凌晨"
        table.cell(i, 2).text = u"{:.1f}".format(self.data["kurtosis0"])
        table.cell(i, 3).text = u""
        table.cell(i, 4).text = u"优"

        table.cell(i+1, 1).text = u"过车"
        table.cell(i+1, 2).text = u"{:.1f}".format(self.data["kurtosis2"])
        table.cell(i+1, 3).text = u""
        table.cell(i+1, 4).text = u"中"

        table.cell(i, 0).merge(table.cell(i+1, 0))

        i+=2
        # 主频
        table.cell(i, 0).text = u"主频"
        table.cell(i, 1).text = u"凌晨"
        table.cell(i, 2).text = u"{:.1f}".format(self.data["mainFreq0"])
        table.cell(i, 3).text = u""
        table.cell(i, 4).text = u"优"

        table.cell(i+1, 1).text = u"过车"
        table.cell(i+1, 2).text = u"{:.1f}".format(self.data["mainFreq2"])
        table.cell(i+1, 3).text = u""
        table.cell(i+1, 4).text = u"中"

        table.cell(i, 0).merge(table.cell(i+1, 0))

        i+=2
        # 主频能量
        table.cell(i, 0).text = u"主频能量"
        table.cell(i, 1).text = u"凌晨"
        table.cell(i, 2).text = u"{:.1f}".format(self.data["ampFreq0"])
        table.cell(i, 3).text = u""
        table.cell(i, 4).text = u"优"

        table.cell(i+1, 1).text = u"过车"
        table.cell(i+1, 2).text =  u"{:.1f}".format(self.data["ampFreq2"])
        table.cell(i+1, 3).text = u""
        table.cell(i+1, 4).text = u"中"

        table.cell(i, 0).merge(table.cell(i+1, 0))

        i+=2
        # 过车响应时长
        table.cell(i, 0).text = u"过车响应时长"
        table.cell(i, 1).text = u"过车"
        table.cell(i, 2).text = u"0.4"
        table.cell(i, 3).text = u""
        table.cell(i, 4).text = u"优"

        p2 = self.document.add_paragraph("表3-1 主要监测指标")
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.style = self.document.styles["TOCHeading"]
        return True


    def _gen_four_paragraph(self):
        p = self.document.add_heading(u"4.区域横向对比", level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        p1 = self.document.add_heading(u"4.1减震区", level=2)
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT

        p1_1 = self.document.add_heading(u"4.1.1凌晨", level=3)
        p1_1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        startpar = self.draw_pics(1)


        p1_2 = self.document.add_heading(u"4.1.2过车", level=3)
        p1_2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        startpar = self.draw_pics(startpar)

        p2 = self.document.add_heading(u"4.2非减震区", level=2)
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT

        p2_1 = self.document.add_heading(u"4.2.1凌晨", level=3)
        p2_1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # add four picture 
        startpar = self.draw_pics(startpar)
        p2_2 = self.document.add_heading(u"4.2.2过车", level=3)
        p2_2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        startpar = self.draw_pics(startpar)


    def draw_pics(self,  startpar):
        filename_zero = self.tmpdir + "/damping_waterfall_zero_%d.png" % startpar
        draw_waterfall(filename_zero)
        self.document.add_picture(filename_zero)
        p_zero = self.document.add_paragraph("图4-%d 瀑布图展示" % startpar)
        p_zero.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_zero.style = self.document.styles["TOCHeading"]

        startpar += 1
        filename_temp = self.tmpdir + "/aserias_compare_%d.png" % startpar
        draw_a_serias(filename_temp)
        self.document.add_picture(filename_temp)
        p_temp = self.document.add_paragraph("图4-%d 各区域A系指标横向变化折线图" % startpar)
        p_temp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_temp.style = self.document.styles["TOCHeading"]
        # add four picture 
        startpar += 1
        filename_temp = self.tmpdir + "/bserias_compare_%d.png" % startpar
        draw_b_serias(filename_temp)
        self.document.add_picture(filename_temp)
        p_temp = self.document.add_paragraph("图4-%d 各区域B系指标横向变化折线图" % startpar)
        p_temp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_temp.style = self.document.styles["TOCHeading"]

        startpar += 1
        filename_temp = self.tmpdir + "/fluction_compare_%d.png" % startpar
        draw_fluctuation(filename_temp)
        self.document.add_picture(filename_temp)
        p_temp = self.document.add_paragraph("图4-%d 各区域波动程度横向变化折线图" % startpar)
        p_temp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_temp.style = self.document.styles["TOCHeading"]

        startpar += 1
        filename_temp = self.tmpdir + "/freq_compare_%d.png" % startpar
        draw_freq(filename_temp)
        self.document.add_picture(filename_temp)
        p_temp = self.document.add_paragraph("图4-%d 各区域传感器监测的主频横向变化折线图" % startpar)
        p_temp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_temp.style = self.document.styles["TOCHeading"]
        startpar += 1
        return startpar

    def _save_report(self, filename="report.docx"):
        print(filename)
        self.document.save(filename)

